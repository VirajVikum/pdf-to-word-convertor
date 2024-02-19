import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import fitz  # PyMuPDF
import pytesseract
import io
import os
import subprocess
import sys
import threading  # Import threading module

class PDFToWordConverterApp:
    def __init__(self, master):
        self.master = master
        master.title("PDF to Word Converter")
        master.geometry("400x395")  # Set the window size
        master.resizable(False, False)

        # Set window background color to lightcyan
        master.configure(bg="black")

        self.label_instruction = tk.Label(master, text="PDF to DOCX convertor", bg="lime",fg="black", width=30, height=2, font=("Arial", 20, "bold"))
        self.label_instruction.pack(pady=0)

        self.browse_button = tk.Button(master,text="Browse", command=self.browse_file, width=20, height=2, bg="lime", fg="black", font=("Arial", 11))
        self.browse_button.pack(pady=40)

        self.label_selected_file = tk.Label(master, text="Browse your PDF File ", bg="black",fg="white", width=40, font=("Arial", 11))
        self.label_selected_file.pack(pady=0)

        self.convert_button = tk.Button(master, text="Convert to Word", command=self.convert_to_word, width=20, height=2, bg="lime", fg="black", font=("Arial", 11))
        self.convert_button.pack(pady=40)

        self.progress_label = tk.Label(master, text="", bg="black",fg="lightgreen", font=("Arial", 11))
        self.progress_label.pack(pady=0)

        self.progress_bar = ttk.Progressbar(master, orient="horizontal", length=400, mode="determinate")
        self.progress_bar.pack(pady=5)

        # Bind close event
        master.protocol("WM_DELETE_WINDOW", self.on_closing)

    def browse_file(self):
        self.reset_progress_bar()
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        filename = os.path.basename(file_path)
        if file_path:
            self.label_selected_file.config(text="Selected File: " + filename)
            self.pdf_path = file_path

    def convert_to_word(self):
        if hasattr(self, 'pdf_path'):
            output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if output_path:
                # Disable the convert button during conversion
                self.convert_button.config(state=tk.DISABLED, font=("Arial", 11))
                self.browse_button.config(state=tk.DISABLED, font=("Arial", 11))

                # Start conversion in a separate thread
                threading.Thread(target=self.pdf_to_word, args=(self.pdf_path, output_path)).start()
        else:
            messagebox.showwarning("File Not Selected", "Please select a PDF file before converting.")

    def reset_progress_bar(self):
        self.progress_label.config(text="")
        self.progress_bar["value"] = 0
        self.master.update_idletasks()

    def pdf_to_word(self, pdf_path, output_path, margin_inches=1.0, font_size=12):
        doc = Document()

        # Set Word document margins
        section = doc.sections[0]
        section.left_margin = Pt(margin_inches * 72)  # Convert inches to points
        section.right_margin = Pt(margin_inches * 72)
        section.top_margin = Pt(margin_inches * 72)
        section.bottom_margin = Pt(margin_inches * 72)

        images = self.pdf_to_images(pdf_path)

        total_images = len(images)
        for i, image in enumerate(images):
            if i > 0:
                doc.add_page_break()  # Start a new page after processing the first page

            image_path = f"temp_image_{i}.png"
            image.save(image_path)
            extracted_text = self.extract_text_from_image(image_path)

            # Add paragraphs with accurate spacing and font size
            for line in extracted_text.split('\n'):
                paragraph = doc.add_paragraph(line)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                paragraph.paragraph_format.space_before = Pt(12)  # Adjust as needed

                if paragraph.runs:
                    run = paragraph.runs[0]
                else:
                    run = paragraph.add_run()

                run.font.size = Pt(font_size)

            # Update progress bar
            progress_percentage = int((i + 1) / total_images * 100)
            self.progress_label.config(text=f"Converting: {progress_percentage}%")
            self.progress_bar["value"] = progress_percentage
            self.master.update_idletasks()

        doc.save(output_path)
        
        # Re-enable the convert button after conversion
        self.convert_button.config(state=tk.NORMAL, text="Convert to Word")
        self.browse_button.config(state=tk.NORMAL)
        messagebox.showinfo("Conversion Complete", f"Word document created at {output_path}")

        # Open the Word document
        self.open_word_document(output_path)
        self.open_pdf_document(self.pdf_path)

        # Clean up: Delete temporary images
        self.cleanup_temp_images()

        # Reset progress bar
        self.reset_progress_bar()

    def open_word_document(self, doc_path):
        if sys.platform.startswith('darwin'):  # macOS
            subprocess.call(('open', doc_path))
        elif os.name == 'nt':  # Windows
            os.startfile(doc_path)
        elif os.name == 'posix':  # Linux
            subprocess.call(('xdg-open', doc_path))

    def open_pdf_document(self, pdf_path):
        if sys.platform.startswith('darwin'):  # macOS
            subprocess.call(('open', pdf_path))
        elif os.name == 'nt':  # Windows
            os.startfile(pdf_path)
        elif os.name == 'posix':  # Linux
            subprocess.call(('xdg-open', pdf_path))

    def cleanup_temp_images(self):
        for i in range(len(self.images)):
            image_path = f"temp_image_{i}.png"
            if os.path.exists(image_path):
                os.remove(image_path)

    def pdf_to_images(self, pdf_path):
        doc = fitz.open(pdf_path)
        self.images = []  # Store images as an instance variable

        for page_num in range(doc.page_count):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                base_image = doc.extract_image(img_info[0])
                image_bytes = base_image["image"]

                image = Image.open(io.BytesIO(image_bytes))
                self.images.append(image)

        return self.images

    def extract_text_from_image(self, image_path):
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang='eng')
        return text

    def on_closing(self):
        if messagebox.askokcancel("Quit", "Do you want to quit?"):
            self.master.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFToWordConverterApp(root)
    root.mainloop()
